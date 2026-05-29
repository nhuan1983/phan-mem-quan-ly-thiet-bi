import streamlit as st
import pandas as pd
import datetime
from io import BytesIO
from docx import Document
import plotly.express as px

# ==========================================
# CẤU HÌNH GIAO DIỆN & TỐI ƯU HÓA MOBILE
# ==========================================
st.set_page_config(page_title="Quản lý Thiết bị KHTN", layout="wide")

hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .mobile-hint {
        background-color: #1E88E5;
        color: white;
        padding: 10px;
        text-align: center;
        font-size: 14px;
        margin-bottom: 15px;
        border-radius: 5px;
        display: none;
    }
    @media only screen and (max-width: 600px) {
        .mobile-hint { display: block; }
    }
    </style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
st.markdown('<div class="mobile-hint">📱 Nhấn vào biểu tượng <b>></b> hoặc <b>☰</b> (góc trái) để mở menu chức năng</div>', unsafe_allow_html=True)

# ==========================================
# CẤU HÌNH KẾT NỐI GOOGLE SHEETS
# ==========================================
USE_CLOUD_DB = False
sh = None

if "gspread_creds" in st.secrets and "spreadsheet_key" in st.secrets:
    try:
        import gspread
        creds_dict = dict(st.secrets["gspread_creds"])
        if "private_key" in creds_dict:
            creds_dict["private_key"] = creds_dict["private_key"].replace("\\n", "\n")
            
        gc = gspread.service_account_from_dict(creds_dict)
        sh = gc.open_by_key(st.secrets["spreadsheet_key"])
        USE_CLOUD_DB = True
    except Exception as e:
        st.sidebar.error(f"⚠️ Lỗi kết nối CSDL Đám mây: {e}")

# --- CÁC HÀM ĐỌC/GHI DỮ LIỆU ĐÃ ĐƯỢC CHUẨN HÓA ĐỂ LƯU CLOUD ---
def load_data(sheet_name, default_df):
    if USE_CLOUD_DB:
        try:
            worksheet = sh.worksheet(sheet_name)
            records = worksheet.get_all_records()
            if not records:
                return default_df
            df = pd.DataFrame(records)
            if sheet_name == 'users' and 'Vai trò' in df.columns:
                df['Vai trò'] = df['Vai trò'].apply(lambda x: [r.strip() for r in str(x).split(',')])
            return df
        except Exception:
            return default_df
    return default_df

def save_data(sheet_name, df_to_save):
    if USE_CLOUD_DB:
        try:
            try:
                worksheet = sh.worksheet(sheet_name)
            except gspread.exceptions.WorksheetNotFound:
                worksheet = sh.add_worksheet(title=sheet_name, rows="100", cols="20")
            
            worksheet.clear()
            df_copy = df_to_save.copy()
            
            if sheet_name == 'users' and 'Vai trò' in df_copy.columns:
                df_copy['Vai trò'] = df_copy['Vai trò'].apply(lambda x: ", ".join(x) if isinstance(x, list) else x)
            
            # Ép kiểu toàn bộ dữ liệu về chuỗi (string) để tránh lỗi định dạng khi lưu lên Cloud
            for col in df_copy.columns:
                df_copy[col] = df_copy[col].astype(str).replace('nan', '')
                
            worksheet.update(values=[df_copy.columns.values.tolist()] + df_copy.values.tolist(), range_name="A1")
        except Exception as e:
            st.error(f"Không thể đồng bộ dữ liệu: {e}")

# ==========================================
# KHỞI TẠO HOẶC TẢI CƠ SỞ DỮ LIỆU
# ==========================================
default_school = pd.DataFrame([{'ten_truong': 'TH&THCS Nam Thượng', 'don_vi_chu_quan': 'Phòng GD&ĐT', 'nam_hoc': '2025-2026'}])
df_school_db = load_data('school_info', default_school)
if 'school_info' not in st.session_state:
    st.session_state.school_info = df_school_db.iloc[0].to_dict()

default_users = pd.DataFrame({
    'Tài khoản': ['admin', 'ht', 'totruong', 'gv01'],
    'Mật khẩu': ['123', '123', '123', '123'],
    'Họ tên': ['Quản trị viên (PHT)', 'Nguyễn Văn A (Hiệu trưởng)', 'Trần Thị B (Tổ trưởng)', 'Lê Văn C (Giáo viên)'],
    'Vai trò': [['Quản trị viên', 'Phó Hiệu trưởng', 'Giáo viên bộ môn'], ['Hiệu trưởng', 'Giáo viên bộ môn'], ['Tổ trưởng chuyên môn', 'Giáo viên bộ môn'], ['Giáo viên bộ môn']]
})
if 'users' not in st.session_state:
    st.session_state.users = load_data('users', default_users)

default_chem = pd.DataFrame({
    'Mã vật tư': ['HC01', 'HC02', 'VL01', 'SH01', 'DC01'],
    'Tên vật tư': ['Axit Sunfuric (H2SO4)', 'Natri Hidroxit (NaOH)', 'Bộ Khúc xạ ánh sáng', 'Tiêu bản tế bào', 'Kính hiển vi quang học'],
    'Phân môn': ['Hóa học', 'Hóa học', 'Vật lý', 'Sinh học', 'Dùng chung'],
    'Số lượng': [500, 1000, 3, 20, 5],
    'Đơn vị': ['ml', 'gam', 'bộ', 'cái', 'cái'],
    'Hạn sử dụng': ['15/06/2026', '10/04/2026', '', '01/01/2027', ''],
    'Tình trạng': ['Tốt', 'Sắp hết hạn', 'Tốt', 'Tốt', 'Tốt']
})
if 'chemicals' not in st.session_state:
    st.session_state.chemicals = load_data('chemicals', default_chem)

if 'bookings' not in st.session_state:
    st.session_state.bookings = load_data('bookings', pd.DataFrame(columns=['Người đăng ký', 'Ngày', 'Buổi', 'Tiết', 'Lớp', 'Môn', 'Thiết bị']))

if 'evaluations' not in st.session_state:
    df_eval_loaded = load_data('evaluations', pd.DataFrame())
    st.session_state.evaluations = df_eval_loaded.to_dict('records') if not df_eval_loaded.empty else []

if 'logged_in' not in st.session_state: st.session_state.logged_in = False
if 'current_user' not in st.session_state: st.session_state.current_user = None

# ==========================================
# GIAO DIỆN ĐĂNG NHẬP (Đã khắc phục lỗi định dạng chuỗi mật khẩu)
# ==========================================
if not st.session_state.logged_in:
    st.markdown("<h1 style='text-align: center; color: #1E88E5;'>HỆ THỐNG QUẢN LÝ THIẾT BỊ DẠY HỌC</h1>", unsafe_allow_html=True)
    st.markdown(f"<h3 style='text-align: center;'>Trường {st.session_state.school_info['ten_truong']}</h3>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        with st.form("login_form"):
            st.write("Vui lòng đăng nhập để tiếp tục")
            username = st.text_input("Tài khoản")
            password = st.text_input("Mật khẩu", type="password")
            if st.form_submit_button("Đăng nhập", use_container_width=True):
                user_match = st.session_state.users[(st.session_state.users['Tài khoản'].astype(str) == str(username)) & (st.session_state.users['Mật khẩu'].astype(str) == str(password))]
                if not user_match.empty:
                    st.session_state.logged_in = True
                    st.session_state.current_user = user_match.iloc[0].to_dict()
                    st.rerun()
                else: 
                    st.error("Sai tài khoản hoặc mật khẩu!")
    st.stop()

# ==========================================
# THANH ĐIỀU HƯỚNG BÊN TRÁI & MENU
# ==========================================
current_user = st.session_state.current_user
user_roles = current_user['Vai trò']

st.sidebar.title(st.session_state.school_info['ten_truong'])
st.sidebar.caption("🟢 Đã kết nối CSDL đám mây" if USE_CLOUD_DB else "🔵 Chế độ cục bộ (Lưu trữ tạm thời)")
st.sidebar.success(f"👤 Chào, {current_user['Họ tên']}")
st.sidebar.markdown("---")
active_role = st.sidebar.selectbox("🔄 Bạn đang làm việc với tư cách là:", user_roles)
st.sidebar.markdown("---")

menu_options = ["Trang chủ & Cảnh báo", "Quản lý Kho (Vật tư)", "Đăng ký thiết bị", "Đổi mật khẩu"]
if active_role in ["Quản trị viên", "Hiệu trưởng", "Phó Hiệu trưởng", "Tổ trưởng chuyên môn"]:
    menu_options.extend(["Đánh giá chuyên môn", "Xuất báo cáo (.docx)"])
if active_role == "Quản trị viên":
    menu_options.insert(0, "Quản lý Hệ thống (Admin)")

menu = st.sidebar.radio("📌 Chọn chức năng:", menu_options)
if st.sidebar.button("Đăng xuất"):
    st.session_state.logged_in = False
    st.session_state.current_user = None
    st.rerun()

# ==========================================
# 1. QUẢN LÝ HỆ THỐNG (Admin)
# ==========================================
if menu == "Quản lý Hệ thống (Admin)":
    if active_role != "Quản trị viên":
        st.error("⚠️ Bạn không có quyền truy cập khu vực này!")
        st.stop()
        
    st.header("⚙️ Quản lý Hệ thống & Cấu hình Đơn vị")
    
    st.subheader("1. 🏫 Cấu hình thông tin Trường học")
    with st.form("school_config_form"):
        sc_col1, sc_col2, sc_col3 = st.columns(3)
        edit_ten_truong = sc_col1.text_input("Tên trường", value=st.session_state.school_info['ten_truong'])
        edit_chu_quan = sc_col2.text_input("Đơn vị chủ quản", value=st.session_state.school_info['don_vi_chu_quan'])
        edit_nam_hoc = sc_col3.text_input("Năm học", value=st.session_state.school_info['nam_hoc'])
        if st.form_submit_button("💾 Lưu cấu hình"):
            st.session_state.school_info['ten_truong'] = edit_ten_truong
            st.session_state.school_info['don_vi_chu_quan'] = edit_chu_quan
            st.session_state.school_info['nam_hoc'] = edit_nam_hoc
            save_data('school_info', pd.DataFrame([st.session_state.school_info]))
            st.success("Đã lưu và đồng bộ cấu hình!")
            st.rerun()

    st.markdown("---")
    st.subheader("2. Quản lý tài khoản hiện tại")
    df_display = st.session_state.users.copy()
    df_display['Vai trò'] = df_display['Vai trò'].apply(lambda x: ", ".join(x) if isinstance(x, list) else x)
    st.dataframe(df_display, use_container_width=True)

    st.subheader("3. Cấp tài khoản mới (Thủ công)")
    with st.form("add_user_form"):
        col1, col2 = st.columns(2)
        with col1:
            new_acc = st.text_input("Tên đăng nhập")
            new_pwd = st.text_input("Mật khẩu")
        with col2:
            new_name = st.text_input("Họ và tên")
            new_roles = st.multiselect("Phân quyền nhóm người dùng", ["Giáo viên bộ môn", "Tổ trưởng chuyên môn", "Phó Hiệu trưởng", "Hiệu trưởng", "Quản trị viên"])
        if st.form_submit_button("Tạo tài khoản"):
            if new_acc and new_pwd and new_name and len(new_roles) > 0:
                new_row = pd.DataFrame([{'Tài khoản': new_acc, 'Mật khẩu': new_pwd, 'Họ tên': new_name, 'Vai trò': new_roles}])
                st.session_state.users = pd.concat([st.session_state.users, new_row], ignore_index=True)
                save_data('users', st.session_state.users)
                st.success("Đã cấp tài khoản thành công!")
                st.rerun()

    st.markdown("---")
    st.subheader("4. 📥 Nhập tài khoản từ file Excel")
    df_mau_tk = pd.DataFrame({'Tài khoản': ['gv_toan01'], 'Mật khẩu': ['123'], 'Họ tên': ['Trần Thị D'], 'Vai trò': ['Giáo viên bộ môn']})
    output_tk = BytesIO()
    with pd.ExcelWriter(output_tk, engine='openpyxl') as writer:
        df_mau_tk.to_excel(writer, index=False)
    st.download_button("⬇️ Tải file Excel mẫu", data=output_tk.getvalue(), file_name="Mau_Tai_Khoan.xlsx")
    
    uploaded_users = st.file_uploader("Chọn file Excel tài khoản đã điền", type=['xlsx', 'xls'])
    if uploaded_users is not None and st.button("Tiến hành nhập dữ liệu"):
        try:
            df_new = pd.read_excel(uploaded_users)
            df_new['Vai trò'] = df_new['Vai trò'].astype(str).apply(lambda x: [r.strip() for r in x.split(',')])
            df_new['Tài khoản'] = df_new['Tài khoản'].astype(str)
            df_new['Mật khẩu'] = df_new['Mật khẩu'].astype(str)
            st.session_state.users = pd.concat([st.session_state.users, df_new], ignore_index=True)
            save_data('users', st.session_state.users)
            st.success("Nhập thành công tài khoản hàng loạt!")
            st.rerun()
        except Exception as e:
            st.error(f"Lỗi: {e}")

# ==========================================
# 2. TRANG CHỦ & CẢNH BÁO
# ==========================================
elif menu == "Trang chủ & Cảnh báo":
    st.header("📊 Bảng điều khiển Tổng quan")
    today = pd.Timestamp.today().normalize()
    df_chem_check = st.session_state.chemicals.copy()
    df_chem_check['Hạn sử dụng'] = pd.to_datetime(df_chem_check['Hạn sử dụng'], dayfirst=True, errors='coerce')
    df_exp = df_chem_check.dropna(subset=['Hạn sử dụng'])
    
    df_warning = pd.DataFrame()
    if not df_exp.empty:
        df_warning = df_exp[(df_exp['Hạn sử dụng'] - today).dt.days <= 30]
        df_warning['Hạn sử dụng'] = df_warning['Hạn sử dụng'].dt.strftime('%d/%m/%Y')
        
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("📦 Tổng mã vật tư", len(st.session_state.chemicals))
    col2.metric("📝 Tổng lượt mượn", len(st.session_state.bookings))
    col3.metric("📋 Hồ sơ chuyên môn", len(st.session_state.evaluations))
    col4.metric("⚠️ Vật tư sắp hết hạn", len(df_warning))
    st.markdown("---")
    
    chart_col1, chart_col2 = st.columns(2)
    with chart_col1:
        st.subheader("Cơ cấu Vật tư")
        if not st.session_state.chemicals.empty:
            df_mon = st.session_state.chemicals['Phân môn'].value_counts().reset_index()
            df_mon.columns = ['Phân môn', 'Số lượng']
            fig1 = px.pie(df_mon, values='Số lượng', names='Phân môn', hole=0.4, color_discrete_sequence=px.colors.qualitative.Pastel)
            st.plotly_chart(fig1, use_container_width=True)
            
    with chart_col2:
        st.subheader("Lượt mượn theo Môn")
        if not st.session_state.bookings.empty:
            df_muon = st.session_state.bookings['Môn'].value_counts().reset_index()
            df_muon.columns = ['Môn học', 'Số lượt']
            fig2 = px.bar(df_muon, x='Môn học', y='Số lượt', color='Môn học', text='Số lượt', color_discrete_sequence=px.colors.qualitative.Set2)
            st.plotly_chart(fig2, use_container_width=True)
            
    st.markdown("---")
    st.subheader("⚠️ Cảnh báo An toàn")
    if not df_warning.empty:
        with st.expander(f"🚨 Có {len(df_warning)} vật tư sắp hoặc đã hết hạn! (Nhấn xem)", expanded=False):
            st.dataframe(df_warning, use_container_width=True)
    else:
        st.success("Tất cả hóa chất/tiêu bản đều trong hạn sử dụng an toàn.")

# ==========================================
# 3. QUẢN LÝ KHO (VẬT TƯ) - CHẮC CHẮN LƯU CLOUD
# ==========================================
elif menu == "Quản lý Kho (Vật tư)":
    st.header("📦 Quản lý Kho Thiết bị & Hóa chất")
    st.dataframe(st.session_state.chemicals, use_container_width=True)
    
    if active_role in ["Quản trị viên", "Hiệu trưởng", "Phó Hiệu trưởng", "Tổ trưởng chuyên môn"]:
        st.markdown("---")
        st.subheader("1. ➕ Bổ sung vật tư mới")
        with st.form("add_chem_form"):
            c1, c2, c3, c4, c5 = st.columns(5)
            ma_vt = c1.text_input("Mã vật tư")
            ten_vt = c2.text_input("Tên vật tư")
            phan_mon = c3.selectbox("Phân môn", ["Vật lý", "Hóa học", "Sinh học", "Dùng chung"])
            don_vi = c4.selectbox("Đơn vị", ["cái", "bộ", "gam", "ml", "gói", "cuộn", "chiếc"])
            so_luong = c5.number_input("Số lượng", min_value=1, value=1)
            
            c6, c7 = st.columns(2)
            han_su_dung = c6.date_input("Hạn sử dụng", value=None, format="DD/MM/YYYY")
            tinh_trang = c7.selectbox("Tình trạng", ["Tốt", "Cần sửa chữa", "Đang đặt mua"])
            
            if st.form_submit_button("Lưu vào kho"):
                if ma_vt and ten_vt and don_vi:
                    han_str = han_su_dung.strftime('%d/%m/%Y') if han_su_dung else ""
                    new_item = pd.DataFrame([{
                        'Mã vật tư': ma_vt, 
                        'Tên vật tư': ten_vt, 
                        'Phân môn': phan_mon, 
                        'Số lượng': int(so_luong), 
                        'Đơn vị': don_vi, 
                        'Hạn sử dụng': han_str, 
                        'Tình trạng': tinh_trang
                    }])
                    st.session_state.chemicals = pd.concat([st.session_state.chemicals, new_item], ignore_index=True)
                    # GỌI LỆNH LƯU LÊN GOOGLE SHEETS
                    save_data('chemicals', st.session_state.chemicals)
                    st.success("✅ Đã bổ sung thiết bị và đồng bộ lên kho đám mây!")
                    st.rerun()
                else:
                    st.warning("⚠️ Vui lòng nhập đủ Mã vật tư, Tên vật tư và Đơn vị!")

        st.markdown("---")
        st.subheader("2. 📥 Nhập hàng loạt vật tư từ file Excel")
        df_mau_vt = pd.DataFrame({'Mã vật tư': ['VL02'], 'Tên vật tư': ['Ampe kế'], 'Phân môn': ['Vật lý'], 'Số lượng': [15], 'Đơn vị': ['cái'], 'Hạn sử dụng': ['15/09/2026'], 'Tình trạng': ['Tốt']})
        output_vt = BytesIO()
        with pd.ExcelWriter(output_vt, engine='openpyxl') as writer:
            df_mau_vt.to_excel(writer, index=False)
        st.download_button("⬇️ Tải file Excel mẫu", data=output_vt.getvalue(), file_name="Mau_Nhap_Vat_Tu.xlsx")
        
        uploaded_chem = st.file_uploader("Chọn file Excel thiết bị đã điền", type=['xlsx', 'xls'])
        if uploaded_chem is not None and st.button("Tiến hành nhập dữ liệu"):
            try:
                df_new = pd.read_excel(uploaded_chem)
                for col in df_new.columns:
                    if df_new[col].dtype == 'object' or 'Hạn sử dụng' in col:
                        df_new[col] = df_new[col].astype(str).replace('nan', '')
                st.session_state.chemicals = pd.concat([st.session_state.chemicals, df_new], ignore_index=True)
                # GỌI LỆNH LƯU LÊN GOOGLE SHEETS
                save_data('chemicals', st.session_state.chemicals)
                st.success("✅ Đã đồng bộ dữ liệu từ Excel lên đám mây thành công!")
                st.rerun()
            except Exception as e:
                st.error(f"Lỗi đọc file: {e}")

        st.markdown("---")
        st.subheader("3. ✏️ Sửa hoặc ❌ Xóa thiết bị")
        item_list = st.session_state.chemicals['Mã vật tư'].tolist()
        selected_item_code = st.selectbox("Chọn Mã vật tư:", ["-- Chọn vật tư --"] + item_list)
        if selected_item_code != "-- Chọn vật tư --":
            item_data = st.session_state.chemicals[st.session_state.chemicals['Mã vật tư'] == selected_item_code].iloc[0]
            with st.form("edit_delete_item_form"):
                edit_ten = st.text_input("Tên vật tư", value=item_data['Tên vật tư'])
                col_e1, col_e2, col_e3 = st.columns(3)
                phan_mon_list = ["Vật lý", "Hóa học", "Sinh học", "Dùng chung"]
                current_mon = item_data['Phân môn'] if item_data['Phân môn'] in phan_mon_list else "Dùng chung"
                edit_mon = col_e1.selectbox("Phân môn", phan_mon_list, index=phan_mon_list.index(current_mon))
                
                don_vi_list = ["cái", "bộ", "gam", "ml", "gói", "cuộn", "chiếc"]
                current_dv = item_data['Đơn vị'] if 'Đơn vị' in item_data and item_data['Đơn vị'] in don_vi_list else "cái"
                edit_dv = col_e2.selectbox("Đơn vị", don_vi_list, index=don_vi_list.index(current_dv) if current_dv in don_vi_list else 0)
                
                default_sl = int(item_data['Số lượng']) if 'Số lượng' in item_data and pd.notnull(item_data['Số lượng']) else 1
                edit_sl = col_e3.number_input("Số lượng hiện có", min_value=0, value=default_sl)
                
                edit_tt = st.selectbox("Tình trạng", ["Tốt", "Cần sửa chữa", "Đang đặt mua"], index=["Tốt", "Cần sửa chữa", "Đang đặt mua"].index(item_data['Tình trạng']) if item_data['Tình trạng'] in ["Tốt", "Cần sửa chữa", "Đang đặt mua"] else 0)
                
                ic_col1, ic_col2 = st.columns(2)
                if ic_col1.form_submit_button("💾 Lưu thay đổi"):
                    idx = st.session_state.chemicals[st.session_state.chemicals['Mã vật tư'] == selected_item_code].index[0]
                    st.session_state.chemicals.at[idx, 'Tên vật tư'] = edit_ten
                    st.session_state.chemicals.at[idx, 'Phân môn'] = edit_mon
                    st.session_state.chemicals.at[idx, 'Đơn vị'] = edit_dv
                    st.session_state.chemicals.at[idx, 'Số lượng'] = int(edit_sl)
                    st.session_state.chemicals.at[idx, 'Tình trạng'] = edit_tt
                    # GỌI LỆNH LƯU LÊN GOOGLE SHEETS
                    save_data('chemicals', st.session_state.chemicals)
                    st.success("✅ Đã lưu cập nhật lên đám mây!")
                    st.rerun()
                if ic_col2.form_submit_button("❌ XÓA THIẾT BỊ NÀY"):
                    st.markdown("---")
        st.subheader("4. 🗑️ Xóa hàng loạt thiết bị/vật tư")
        st.info("💡 Tính năng này giúp thanh lý nhanh chóng nhiều hóa chất, vật tư đã hết hạn hoặc hư hỏng cùng lúc.")
        
        # 1. Tạo danh sách hiển thị gồm [Mã vật tư] - [Tên vật tư] - [Tình trạng] để dễ chọn
        item_options = [f"{row['Mã vật tư']} - {row['Tên vật tư']} ({row['Tình trạng']})" for idx, row in st.session_state.chemicals.iterrows()]
        
        # 2. Sử dụng multiselect để cho phép chọn nhiều thiết bị
        items_to_delete = st.multiselect("Chọn các vật tư cần thanh lý/xóa:", item_options)
        
        if st.button("❌ Xác nhận xóa hàng loạt", type="primary"):
            if items_to_delete:
                # 3. Tách lấy phần 'Mã vật tư' từ chuỗi văn bản đã chọn
                ma_vt_to_delete = [item.split(" - ")[0] for item in items_to_delete]
                
                # 4. Cập nhật lại kho (giữ lại những vật tư KHÔNG nằm trong danh sách cần xóa)
                st.session_state.chemicals = st.session_state.chemicals[~st.session_state.chemicals['Mã vật tư'].isin(ma_vt_to_delete)].reset_index(drop=True)
                
                # 5. Gọi lệnh đồng bộ lên Google Sheets đám mây
                save_data('chemicals', st.session_state.chemicals)
                
                st.success(f"✅ Đã xóa thành công {len(ma_vt_to_delete)} thiết bị khỏi hệ thống đám mây!")
                st.rerun()
            else:
                st.warning("⚠️ Vui lòng chọn ít nhất một vật tư để thực hiện lệnh xóa.")
                st.session_state.chemicals = st.session_state.chemicals[st.session_state.chemicals['Mã vật tư'] != selected_item_code].reset_index(drop=True)
                    # GỌI LỆNH LƯU LÊN GOOGLE SHEETS
                save_data('chemicals', st.session_state.chemicals)
                st.success("✅ Đã xóa thiết bị khỏi hệ thống đám mây!")
                st.rerun()

# ==========================================
# 4. ĐĂNG KÝ THIẾT BỊ 
# ==========================================
elif menu == "Đăng ký thiết bị":
    st.header("📝 Đăng ký sử dụng phòng bộ môn")
    
    st.subheader("1. Lịch đăng ký toàn trường")
    if not st.session_state.bookings.empty:
        st.dataframe(st.session_state.bookings, use_container_width=True)
    else:
        st.info("Chưa có lịch đăng ký nào từ các giáo viên.")
        
    st.markdown("---")
    st.subheader("2. ➕ Tạo lịch đăng ký mới")
    
    col1, col2 = st.columns(2)
    with col1:
        date = st.date_input("Ngày dạy thực hành", format="DD/MM/YYYY")
        buoi = st.selectbox("Buổi dạy", ["Sáng", "Chiều"])
        period = st.selectbox("Tiết học giảng dạy", [1, 2, 3, 4, 5])
        lop = st.text_input("Lớp (VD: 9A)")
        subject_filter = st.selectbox("Phân môn", ["Vật lý", "Hóa học", "Sinh học"])
        
    with col2:
        filtered_equip = st.session_state.chemicals[
            st.session_state.chemicals['Phân môn'].isin([subject_filter, "Dùng chung"])
        ]['Tên vật tư'].tolist()
        equipment = st.multiselect("Chọn thiết bị cần mượn", filtered_equip)
        
    if st.button("Xác nhận đăng ký", type="primary"):
        if lop and equipment:
            date_str = date.strftime('%d/%m/%Y')
            new_book = pd.DataFrame([{'Người đăng ký': current_user['Họ tên'], 'Ngày': date_str, 'Buổi': buoi, 'Tiết': period, 'Lớp': lop, 'Môn': subject_filter, 'Thiết bị': ", ".join(equipment)}])
            st.session_state.bookings = pd.concat([st.session_state.bookings, new_book], ignore_index=True)
            save_data('bookings', st.session_state.bookings)
            st.success("Lịch đăng ký đã được lưu thành công!")
            st.rerun()
        else:
            st.warning("⚠️ Vui lòng điền tên Lớp và chọn ít nhất 1 thiết bị cần mượn!")

    st.markdown("---")
    st.subheader("3. ✏️ Quản lý lịch cá nhân")
    st.info("Giáo viên chỉ có thể xem và điều chỉnh các lịch đăng ký do chính mình tạo ra.")
    
    user_bookings = st.session_state.bookings[st.session_state.bookings['Người đăng ký'] == current_user['Họ tên']]
    
    if not user_bookings.empty:
        booking_options = [f"Lớp {row['Lớp']} - Tiết {row['Tiết']} - Môn {row['Môn']} (Ngày {row['Ngày']})" for idx, row in user_bookings.iterrows()]
        selected_booking_str = st.selectbox("Chọn lịch Thầy/Cô muốn điều chỉnh:", ["-- Chọn lịch --"] + booking_options)
        
        if selected_booking_str != "-- Chọn lịch --":
            actual_idx = user_bookings.index[booking_options.index(selected_booking_str) - 1]
            booking_data = st.session_state.bookings.iloc[actual_idx]
            
            current_subject_edit = booking_data['Môn']
            edit_filtered_equip = st.session_state.chemicals[
                st.session_state.chemicals['Phân môn'].isin([current_subject_edit, "Dùng chung"])
            ]['Tên vật tư'].tolist()
            
            with st.form("edit_booking_form"):
                e_col1, e_col2 = st.columns(2)
                with e_col1:
                    try:
                        parsed_date = pd.to_datetime(booking_data['Ngày'], dayfirst=True).date()
                    except:
                        parsed_date = datetime.date.today()
                        
                    e_date = st.date_input("Sửa ngày dạy", value=parsed_date, format="DD/MM/YYYY")
                    e_buoi = st.selectbox("Sửa buổi dạy", ["Sáng", "Chiều"], index=["Sáng", "Chiều"].index(booking_data['Buổi']))
                    e_period = st.selectbox("Sửa tiết học", [1, 2, 3, 4, 5], index=[1, 2, 3, 4, 5].index(int(booking_data['Tiết'])))
                    e_lop = st.text_input("Sửa Lớp", value=booking_data['Lớp'])
                    st.info(f"Đang sửa thiết bị cho phân môn: **{current_subject_edit}**")
                    
                with e_col2:
                    current_equip = [x.strip() for x in str(booking_data['Thiết bị']).split(',')] if booking_data['Thiết bị'] else []
                    safe_current_equip = [x for x in current_equip if x in edit_filtered_equip]
                    e_equipment = st.multiselect("Chỉnh sửa thiết bị cần mượn", edit_filtered_equip, default=safe_current_equip)
                    
                btn_c1, btn_c2 = st.columns(2)
                if btn_c1.form_submit_button("💾 Lưu thay đổi lịch"):
                    st.session_state.bookings.at[actual_idx, 'Ngày'] = e_date.strftime('%d/%m/%Y')
                    st.session_state.bookings.at[actual_idx, 'Buổi'] = e_buoi
                    st.session_state.bookings.at[actual_idx, 'Tiết'] = e_period
                    st.session_state.bookings.at[actual_idx, 'Lớp'] = e_lop
                    st.session_state.bookings.at[actual_idx, 'Thiết bị'] = ", ".join(e_equipment)
                    save_data('bookings', st.session_state.bookings)
                    st.success("✅ Đã cập nhật lại lịch đăng ký!")
                    st.rerun()
                    
                if btn_c2.form_submit_button("❌ Xóa/Hủy lịch này"):
                    st.session_state.bookings = st.session_state.bookings.drop(actual_idx).reset_index(drop=True)
                    save_data('bookings', st.session_state.bookings)
                    st.success("✅ Đã hủy lịch đăng ký thành công!")
                    st.rerun()
    else:
        st.info("Hiện tại Thầy/Cô chưa có lịch đăng ký nào trên hệ thống.")

# ==========================================
# 5. ĐÁNH GIÁ CHUYÊN MÔN
# ==========================================
elif menu == "Đánh giá chuyên môn":
    if active_role not in ["Quản trị viên", "Hiệu trưởng", "Phó Hiệu trưởng", "Tổ trưởng chuyên môn"]:
        st.error("⚠️ Bạn không có quyền truy cập khu vực này!")
        st.stop()
        
    st.header("📋 Đánh giá năng lực thực hành")
    list_gv = st.session_state.users['Họ tên'].tolist()
    target_gv = st.selectbox("1. Chọn Giáo viên để đánh giá:", ["-- Chọn người được đánh giá --"] + list_gv)
    
    if target_gv != "-- Chọn người được đánh giá --":
        gv_bookings = st.session_state.bookings[st.session_state.bookings['Người đăng ký'] == target_gv]
        if gv_bookings.empty:
            st.warning("Giáo viên này chưa đăng ký tiết dạy.")
        else:
            booking_options = [f"Ngày {row['Ngày']} - Buổi {row['Buổi']} - Tiết {row['Tiết']} - Lớp {row['Lớp']} - Môn {row['Môn']}" for idx, row in gv_bookings.iterrows()]
            target_tiet = st.selectbox("2. Chọn Tiết dự giờ:", booking_options)
            
            st.markdown("### 3. Chấm điểm Rubric")
            c1 = st.slider("1. Chuẩn bị thiết bị vật tư (Tối đa 20đ)", 0, 20, 15)
            c2 = st.slider("2. Đảm bảo an toàn PTN (Tối đa 30đ)", 0, 30, 25)
            c3 = st.slider("3. Hướng dẫn HS thao tác (Tối đa 30đ)", 0, 30, 25)
            c4 = st.slider("4. Đánh giá kết quả & Liên hệ PISA (Tối đa 20đ)", 0, 20, 15)
            
            total = c1 + c2 + c3 + c4
            st.markdown(f"**Tổng điểm:** {total}/100")
            rank = "Tốt" if total >= 85 else "Khá" if total >= 65 else "Đạt" if total >= 50 else "Chưa đạt"
            comment = st.text_area("Nhận xét:")
            
            if st.button("Lưu hồ sơ đánh giá"):
                record = {
                    "Người được đánh giá": target_gv, "Tiết dạy": target_tiet, "Người đánh giá": current_user['Họ tên'],
                    "Chức vụ người đánh giá": active_role, "Tổng điểm": total, "Xếp loại": rank, "Nhận xét": comment,
                    "Ngày chấm": datetime.date.today().strftime("%d/%m/%Y")
                }
                st.session_state.evaluations.append(record)
                save_data('evaluations', pd.DataFrame(st.session_state.evaluations))
                st.success("Hồ sơ đã được lưu lên đám mây!")

# ==========================================
# 6. XUẤT BÁO CÁO TỰ ĐỘNG
# ==========================================
elif menu == "Xuất báo cáo (.docx)":
    if active_role not in ["Quản trị viên", "Hiệu trưởng", "Phó Hiệu trưởng", "Tổ trưởng chuyên môn"]:
        st.error("⚠️ Bạn không có quyền truy cập khu vực này!")
        st.stop()
        
    st.header("🖨️ Kết xuất hồ sơ minh chứng")
    if len(st.session_state.evaluations) == 0:
        st.info("Chưa có hồ sơ đánh giá.")
    else:
        st.dataframe(pd.DataFrame(st.session_state.evaluations), use_container_width=True)
        selected_idx = st.selectbox("Chọn hồ sơ:", range(len(st.session_state.evaluations)), format_func=lambda x: f"Đánh giá: {st.session_state.evaluations[x]['Người được đánh giá']} ({st.session_state.evaluations[x]['Tiết dạy']})")
        target_record = st.session_state.evaluations[selected_idx]
        
        def create_docx(data, school_info):
            doc = Document()
            doc.add_heading(school_info['don_vi_chu_quan'].upper(), 1)
            doc.add_heading(f"TRƯỜNG {school_info['ten_truong'].upper()}", 2)
            doc.add_paragraph(f"Năm học: {school_info['nam_hoc']}")
            doc.add_paragraph("-----------------------------------")
            doc.add_heading('PHIẾU ĐÁNH GIÁ NĂNG LỰC THỰC HÀNH', 0)
            doc.add_paragraph(f"Ngày lập: {data['Ngày chấm']}")
            doc.add_paragraph(f"Người đánh giá: {data['Người đánh giá']} ({data['Chức vụ người đánh giá']})")
            doc.add_heading('Thông tin tiết dạy', level=1)
            doc.add_paragraph(f"- Giáo viên: {data['Người được đánh giá']}")
            doc.add_paragraph(f"- Lịch học: {data['Tiết dạy']}")
            doc.add_heading('Kết quả', level=1)
            doc.add_paragraph(f"- Điểm: {data['Tổng điểm']}/100")
            doc.add_paragraph(f"- Xếp loại: {data['Xếp loại']}")
            doc.add_heading('Nhận xét', level=1)
            doc.add_paragraph(data['Nhận xét'])
            bio = BytesIO()
            doc.save(bio)
            return bio.getvalue()

        st.download_button("📄 Tải xuống văn bản (.docx)", data=create_docx(target_record, st.session_state.school_info), file_name=f"PhieuDanhGia_{target_record['Người được đánh giá']}.docx")

# ==========================================
# 7. ĐỔI MẬT KHẨU CÁ NHÂN
# ==========================================
elif menu == "Đổi mật khẩu":
    st.header("🔑 Đổi mật khẩu cá nhân")
    st.info("Lưu ý: Mật khẩu mới sẽ được đồng bộ ngay lập tức.")
    with st.form("change_pass_form"):
        old_pass = st.text_input("Mật khẩu hiện tại", type="password")
        new_pass = st.text_input("Mật khẩu mới", type="password")
        re_new_pass = st.text_input("Nhập lại mật khẩu mới", type="password")
        
        if st.form_submit_button("Xác nhận đổi"):
            if str(old_pass) == str(current_user['Mật khẩu']):
                if str(new_pass) == str(re_new_pass):
                    idx = st.session_state.users[st.session_state.users['Tài khoản'] == current_user['Tài khoản']].index[0]
                    st.session_state.users.at[idx, 'Mật khẩu'] = str(new_pass)
                    save_data('users', st.session_state.users)
                    st.session_state.current_user['Mật khẩu'] = str(new_pass)
                    st.success("✅ Đổi mật khẩu thành công!")
                else:
                    st.error("❌ Mật khẩu nhập lại không khớp!")
            else:
                st.error("❌ Mật khẩu hiện tại không đúng!")
